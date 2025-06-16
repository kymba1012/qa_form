import streamlit as st
import pandas as pd
from docx import Document
from io import BytesIO
import zipfile

# Page setup
st.set_page_config(page_title="Cable Form Generator", layout="centered")
st.title("📄 Cable Form Generator")
st.write("Upload your Excel schedule and Word template to generate custom documents.")

# Upload files
uploaded_excel = st.file_uploader("📊 Upload Excel file", type=["xlsx"])
uploaded_template = st.file_uploader("📄 Upload Word template", type=["docx"])

# Placeholder mapping (same as before)
placeholder_map = {
    "<<Client>>": "Client",
    "<<Project>>": "Project",
    "<<Job>>": "Job",
    "<<Location>>": "Location",
    "<<Area>>": "Area",
    "<<Cable Tag>>": "Cable Tag",
    "<<AWG>>": "AWG",
    "<<Cable Type>>": "Cable Type",
    "<<Number of Conductors>>": "# of Conductors",
    "<<Operating Voltage>>": "Operating Voltage",
    "<<Rated Voltage>>": "Insul Voltage",
    "<<Source>>": "Source",
    "<<Destination>>": "Destination",
    "<<Source Torque>>": "Source Torque Value",
    "<<Destination Torque>>": "Destination Torque Value",
    "<<Test Equipment 1 Make>>": "Test Equipment 1 Make",
    "<<Test Equipment 1 Model>>": "Test Equipment 1 Model",
    "<<Test Equipment 1 Serial>>": "Test Equipment 1 Serial",
    "<<Cal Test Date>>": "Test Equipment 1 Cal Date",
    "<<Test Equipment 1 Calibration Due Date>>": "Test Equipment 1 Calibration Due Date",
    "<<DATE 1>>": "DATE 1",
    "<<Client Rep>>": "Client Rep",
    "<<Client Rep Date>>": "Client Rep Date",
    "<<CDN Rep>>": "CDN Rep",
    "<<CDN Rep Date>>": "CDN Rep Date",
    "<<Comments>>": "Comments"
}

if uploaded_excel and uploaded_template:
    if st.button("🚀 Generate Documents"):
        df = pd.read_excel(uploaded_excel)
        zip_buffer = BytesIO()

        with zipfile.ZipFile(zip_buffer, "a", zipfile.ZIP_DEFLATED) as zipf:
            for i, row in df.iterrows():
                doc = Document(uploaded_template)
                
                # Replace in paragraphs
                for p in doc.paragraphs:
                    for ph, col in placeholder_map.items():
                        if ph in p.text:
                            p.text = p.text.replace(ph, str(row.get(col, "")))

                # Replace in tables
                for table in doc.tables:
                    for r in table.rows:
                        for cell in r.cells:
                            for ph, col in placeholder_map.items():
                                if ph in cell.text:
                                    cell.text = cell.text.replace(ph, str(row.get(col, "")))
                
                cable_id = row.get("Cable Tag", f"row_{i}")
                buffer = BytesIO()
                doc.save(buffer)
                buffer.seek(0)
                zipf.writestr(f"QA_Form_{cable_id}.docx", buffer.read())

        zip_buffer.seek(0)
        st.success("✅ Documents generated!")
        st.download_button("📥 Download ZIP", zip_buffer, file_name="QA_Forms.zip")

